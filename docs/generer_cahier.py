# -*- coding: utf-8 -*-
"""Génère le cahier des charges (Word) avec les diagrammes insérés."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

BASE = os.path.dirname(os.path.abspath(__file__))
DIAG = os.path.join(BASE, "diagrammes")
VERT = RGBColor(0x2E, 0x7D, 0x32)

doc = Document()

# Styles de base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)


def titre(txt, niveau=1):
    h = doc.add_heading(txt, level=niveau)
    for run in h.runs:
        run.font.color.rgb = VERT
    return h


def para(txt):
    p = doc.add_paragraph(txt)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def puce(txt):
    doc.add_paragraph(txt, style='List Bullet')


# ---------- PAGE DE TITRE ----------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("CAHIER DES CHARGES")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = VERT

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Plateforme de Gestion des Coopératives Agricoles")
r.bold = True
r.font.size = Pt(16)

for _ in range(2):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(
    "Université Virtuelle du Burkina Faso\n"
    "Licence 3 — Mathématiques-Informatique\n"
    "Année académique 2025–2026\n\n"
    "Technologies : Python · Django · PostgreSQL · Bootstrap · MoneyFusion"
).font.size = Pt(12)

doc.add_page_break()

# ---------- 1. INTRODUCTION ----------
titre("1. Introduction", 1)
para(
    "Ce document décrit le cahier des charges de la plateforme de gestion des "
    "coopératives agricoles. L'objectif est de donner aux coopératives un outil "
    "numérique unique pour administrer leurs membres, suivre leurs productions et "
    "vendre leurs récoltes en ligne, le tout dans une interface accessible y compris "
    "depuis un simple téléphone."
)
titre("1.1. Présentation générale du projet", 2)
para(
    "Le projet consiste à concevoir et développer une application web pour les "
    "coopératives agricoles du Burkina Faso. Elle réunit en un seul endroit la gestion "
    "administrative des adhérents, le suivi des récoltes et la commercialisation des "
    "produits, là où ces activités étaient jusqu'ici dispersées entre registres papier "
    "et échanges informels."
)
titre("1.2. Importance de l'agriculture au Burkina Faso", 2)
para(
    "L'agriculture reste le moteur de l'économie burkinabè : elle occupe la majeure "
    "partie de la population active et garantit l'essentiel de la sécurité alimentaire. "
    "Mieux organiser les coopératives et numériser leur gestion constitue donc un levier "
    "direct pour augmenter la productivité et les revenus des agriculteurs."
)
titre("1.3. Contexte et motivation", 2)
para(
    "Beaucoup de coopératives travaillent encore avec des cahiers manuscrits et des "
    "suivis informels. Ces méthodes limitent la traçabilité, compliquent le suivi des "
    "cotisations et freinent l'accès à de nouveaux acheteurs. La plateforme répond "
    "précisément à ce besoin de modernisation."
)

# ---------- 2. CONTEXTE ----------
titre("2. Contexte du projet", 1)
titre("2.1. Situation actuelle", 2)
para(
    "La plupart des coopératives gèrent leurs activités sans outil numérique structuré. "
    "Les listes de membres, les volumes récoltés et les transactions sont notés sur "
    "papier, ce qui rend le suivi long et exposé aux erreurs ou aux pertes d'information."
)
titre("2.2. Problématiques rencontrées", 2)
puce("Gestion des membres : difficulté à tenir à jour les adhésions et les cotisations.")
puce("Suivi des productions : pas d'outil centralisé pour les volumes et les stocks.")
puce("Ventes : faible visibilité commerciale et dépendance aux intermédiaires.")
titre("2.3. Objectifs dans ce contexte", 2)
para(
    "Face à ces constats, le projet dote les coopératives d'un outil qui structure la "
    "gestion des membres, fiabilise le suivi des productions et ouvre un canal de vente "
    "en ligne intégré."
)

# ---------- 3. OBJECTIFS ----------
titre("3. Objectifs du projet", 1)
titre("3.1. Objectifs généraux", 2)
puce("Améliorer la gestion administrative et opérationnelle des coopératives.")
puce("Promouvoir la vente des produits agricoles via un canal numérique accessible.")
titre("3.2. Objectifs spécifiques", 2)
puce("Mettre en place un module de gestion des membres (inscription, suivi, cotisations).")
puce("Mettre en place un module de suivi des productions (récoltes, stocks).")
puce("Développer un module e-commerce (catalogue, panier, commandes, paiement).")
puce("Proposer un module de formation en ligne pour les agriculteurs.")

# ---------- 4. ANALYSE DES BESOINS ----------
titre("4. Analyse des besoins", 1)
titre("4.1. Besoins des coopératives", 2)
para(
    "L'analyse fait ressortir le besoin d'un système centralisé capable de gérer en même "
    "temps les membres, les productions, les ventes et la formation, tout en restant simple "
    "pour des utilisateurs dont l'aisance numérique varie beaucoup d'une personne à l'autre."
)
titre("4.2. Utilisateurs cibles", 2)
puce("Agriculteurs (membres) : producteurs et bénéficiaires des formations.")
puce("Administrateurs de coopératives : gestion des membres, des productions et des ventes.")
puce("Acheteurs : particuliers ou structures souhaitant acheter en ligne.")
puce("Formateurs et livreurs : création de contenus pédagogiques et transport des commandes.")
titre("4.3. Environnement actuel", 2)
para(
    "Les outils en place se limitent à des registres papier, des tableurs isolés et des "
    "échanges par téléphone. Il en résulte un manque de traçabilité et de visibilité que "
    "la plateforme vient combler."
)

# ---------- 5. CONCEPTION ET ARCHITECTURE ----------
titre("5. Conception et architecture", 1)
titre("5.1. Choix des technologies", 2)
para(
    "La plateforme repose sur Python et le framework Django pour la partie serveur, "
    "Bootstrap pour l'interface, et PostgreSQL comme système de gestion de base de données. "
    "Le paiement Mobile Money est assuré par l'intégration de la passerelle MoneyFusion. "
    "Ces choix privilégient la robustesse, la sécurité et la rapidité de développement."
)
titre("5.2. Architecture du système", 2)
para(
    "L'application suit le modèle MVT de Django (Modèle-Vue-Template), proche du MVC. Le "
    "code est découpé en applications autonomes — users, productions, sales et formation — "
    "ce qui facilite la maintenance et l'évolution. Le navigateur communique avec les vues "
    "Django, qui s'appuient sur les modèles pour lire et écrire dans PostgreSQL ; les "
    "paiements transitent par MoneyFusion."
)
titre("5.3. Modélisation de la base de données", 2)
para(
    "La base s'organise autour de quinze tables : un compte de connexion commun "
    "(utilisateur) décliné en cinq profils (membre, administrateur, acheteur, formateur, "
    "livreur), puis les tables métier des productions, du catalogue, des commandes, des "
    "paiements, des livraisons, des reçus et des formations. Le schéma complet est présenté "
    "en annexe (Figure 3)."
)

# ---------- 6. IMPLEMENTATION ----------
titre("6. Implémentation", 1)
titre("6.1. Étapes de développement", 2)
para(
    "Le développement a été mené de façon itérative : modélisation de la base, création des "
    "comptes et des rôles, exposition d'une API REST sécurisée par jetons JWT, puis "
    "construction de l'interface web et intégration du paiement. Chaque module a été testé "
    "avant d'être relié aux autres."
)
titre("6.2. Fonctionnalités développées", 2)
puce("Gestion des membres : inscription, profils par rôle, suivi des cotisations.")
puce("Suivi des productions : déclaration des récoltes, validation, gestion du stock.")
puce("Vente en ligne : catalogue, panier d'achat, passation et suivi des commandes.")
puce("Paiement Mobile Money et génération automatique des reçus.")
titre("6.3. Modules clés", 2)
para(
    "Le module des membres centralise les informations des adhérents et leur statut. Le "
    "module des productions enregistre les volumes récoltés et alimente le catalogue une "
    "fois la récolte validée. Le module de vente gère le panier, transforme la commande en "
    "paiement, puis déclenche la livraison et le reçu."
)

# ---------- 7. INTERFACE ----------
titre("7. Interface utilisateur", 1)
titre("7.1. Conception de l'interface", 2)
para(
    "L'interface a été pensée pour aller à l'essentiel : une barre de navigation présente "
    "sur toutes les pages, des formulaires courts et des boutons d'action visibles. Un "
    "utilisateur peu habitué au numérique doit comprendre en quelques secondes ce qu'il "
    "peut faire."
)
titre("7.2. Choix de design", 2)
para(
    "Le design s'appuie sur Bootstrap, ce qui garantit un affichage correct aussi bien sur "
    "ordinateur que sur téléphone — un point essentiel en zone rurale où le mobile domine. "
    "La palette verte rappelle l'univers agricole et renforce le sentiment de confiance."
)
titre("7.3. Principales vues", 2)
para(
    "Les vues principales sont la page d'accueil, le catalogue des produits, le panier, le "
    "tableau de bord adapté au rôle de l'utilisateur, ainsi que l'interface "
    "d'administration. Ces écrans illustrent le parcours complet, de la consultation du "
    "catalogue jusqu'à la confirmation de commande."
)

# ---------- 8. PAIEMENTS ----------
titre("8. Gestion des transactions et paiements", 1)
titre("8.1. Transactions en ligne", 2)
para(
    "Le règlement des commandes s'effectue par Mobile Money grâce à l'intégration de la "
    "passerelle MoneyFusion. À la validation du panier, l'acheteur saisit son numéro, est "
    "redirigé vers la page de paiement sécurisée, puis revient sur la plateforme une fois "
    "l'opération confirmée."
)
titre("8.2. Suivi des paiements et reçus", 2)
para(
    "Chaque paiement est vérifié auprès de la passerelle avant d'être marqué comme réglé. "
    "Dès qu'un paiement passe au statut « payé », un reçu numérique est généré "
    "automatiquement et reste consultable par l'acheteur comme par l'administrateur."
)

# ---------- 9. FORMATION ----------
titre("9. Formation des agriculteurs", 1)
titre("9.1. Module de formation", 2)
para(
    "Un module dédié propose des contenus pédagogiques — fiches, vidéos et quiz — destinés "
    "à renforcer les compétences des agriculteurs sur les techniques de production et la "
    "gestion coopérative."
)
titre("9.2. Suivi des formations", 2)
para(
    "Le système conserve, pour chaque membre, la progression dans les modules suivis et le "
    "score obtenu aux quiz, ce qui permet de garder un historique des formations terminées."
)

# ---------- 10. PERFORMANCES ET SECURITE ----------
titre("10. Évaluation des performances et sécurité", 1)
titre("10.1. Performances", 2)
para(
    "Des tests fonctionnels automatisés couvrent l'inscription, la connexion, le contrôle "
    "des accès par rôle et le parcours d'achat. Ils garantissent la stabilité de la "
    "plateforme, y compris lors des périodes de forte activité comme les récoltes."
)
titre("10.2. Sécurité des données et des transactions", 2)
para(
    "Les mots de passe ne sont jamais stockés en clair : ils sont chiffrés par l'algorithme "
    "de hachage de Django. L'accès à l'API est protégé par jetons JWT et par des permissions "
    "spécifiques à chaque rôle. Les paiements transitent par une passerelle sécurisée et "
    "sont systématiquement vérifiés côté serveur."
)
titre("10.3. Résultats des tests", 2)
para(
    "Les tests automatisés passent intégralement. Le parcours d'achat a été vérifié de bout "
    "en bout, de la création du compte jusqu'à la génération du reçu, et les données "
    "correspondantes sont bien persistées dans PostgreSQL."
)

# ---------- 11. CONCLUSION ----------
titre("11. Conclusion", 1)
titre("11.1. Résumé des réalisations", 2)
para(
    "La plateforme répond aux besoins essentiels des coopératives : gestion des membres, "
    "suivi des productions, vente en ligne avec paiement Mobile Money et formation des "
    "agriculteurs."
)
titre("11.2. Bilan des objectifs", 2)
para(
    "Les objectifs généraux et spécifiques fixés au départ ont guidé l'ensemble des choix de "
    "conception et de développement, et se retrouvent dans les modules livrés."
)
titre("11.3. Perspectives", 2)
para(
    "Plusieurs évolutions sont envisageables : statistiques avancées sur les productions et "
    "les ventes, application mobile dédiée, et élargissement des moyens de paiement."
)

# ---------- 12. ANNEXES ----------
doc.add_page_break()
titre("12. Annexes", 1)
titre("12.1. Diagrammes de modélisation", 2)
para(
    "Les diagrammes ci-dessous illustrent respectivement les interactions des acteurs avec "
    "le système, la structure des classes, le schéma relationnel de la base et le "
    "déroulement d'une commande."
)


def figure(fichier, legende, largeur=6.0):
    chemin = os.path.join(DIAG, fichier)
    doc.add_picture(chemin, width=Inches(largeur))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(legende)
    r.italic = True
    r.font.size = Pt(10)
    doc.add_paragraph()


figure("1_cas_utilisation.png", "Figure 1 — Diagramme de cas d'utilisation", 6.0)
figure("2_classes.png", "Figure 2 — Diagramme de classes", 6.2)
figure("4_sequence_commande.png",
       "Figure 4 — Diagramme de séquence : « Passer une commande »", 6.0)

# Le modèle physique est large : on le met sur une page en paysage.
section = doc.add_section()
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
figure("3_modele_physique.png",
       "Figure 3 — Modèle physique de données (schéma relationnel)", 9.0)

# Retour en portrait
section2 = doc.add_section()
section2.orientation = WD_ORIENT.PORTRAIT
section2.page_width, section2.page_height = section2.page_height, section2.page_width

titre("12.2. Code source", 2)
para(
    "Le code source complet de l'application, ainsi que le script SQL de création de la base "
    "de données (schema.sql), sont fournis en complément du présent document."
)

sortie = os.path.join(BASE, "Cahier_des_charges_Cooperatives_Agricoles.docx")
doc.save(sortie)
print("Document généré :", sortie)
